#!/usr/bin/env python3
"""Build the TMDB movie success prediction report document."""

from __future__ import annotations

import math
import re
import shutil
from pathlib import Path
from typing import Iterable, Sequence

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "outputs" / "report"
FIGURE_DIR = REPORT_DIR / "figures"
REPORT_DOCX = REPORT_DIR / "TMDB_영화흥행예측_결과보고서.docx"
REPORT_MD = REPORT_DIR / "TMDB_영화흥행예측_결과보고서.md"

DATA_FILE = ROOT / "data" / "processed" / "movie_features.csv"
RAW_DATA_FILE = ROOT / "data" / "kobis_tmdb_title_matches.csv"
CLASSIFICATION_METRICS = ROOT / "outputs" / "metrics" / "classification_metrics.csv"
CLASSIFICATION_PREDICTIONS = ROOT / "outputs" / "metrics" / "classification_predictions.csv"
REGRESSION_METRICS = ROOT / "outputs" / "metrics" / "regression_metrics.csv"
RANK_CLASSIFICATION_METRICS = ROOT / "outputs" / "metrics" / "success_rank_classification_metrics.csv"
RANK_CLASSIFICATION_PREDICTIONS = ROOT / "outputs" / "metrics" / "success_rank_classification_predictions.csv"
RANK_METRICS = ROOT / "outputs" / "metrics" / "success_rank_strategy_metrics.csv"
IMPORTANCE_FILE = (
    ROOT
    / "outputs"
    / "metrics"
    / "feature_importance"
    / "classification_all_random_forest_is_success_3000000.csv"
)
CONFUSION_MATRIX = (
    ROOT
    / "outputs"
    / "figures"
    / "confusion_matrix_metadata_random_forest_is_success_3000000.png"
)
VISUALIZATION_UPGRADE_DIR = ROOT / "outputs" / "figures" / "visualization_upgrade"
RANK_ORDER = ["D", "C", "B", "A", "S"]
RANK_CODE = {rank: index for index, rank in enumerate(RANK_ORDER)}
RANK_PALETTE = {
    "D": "#8D99AE",
    "C": "#4C78A8",
    "B": "#59A14F",
    "A": "#F28E2B",
    "S": "#D62728",
}

FONT_CANDIDATES = [
    Path("/System/Library/Fonts/AppleSDGothicNeo.ttc"),
    Path("/System/Library/Fonts/Supplemental/AppleGothic.ttf"),
    Path("/Users/ryankimjh/Library/Fonts/NotoSansCJK.ttc"),
]
FONT_PATH = next((path for path in FONT_CANDIDATES if path.exists()), None)

COLORS = {
    "ink": (38, 45, 56),
    "muted": (102, 112, 133),
    "grid": (220, 226, 235),
    "blue": (45, 114, 210),
    "green": (57, 151, 105),
    "orange": (238, 132, 67),
    "red": (207, 85, 85),
    "purple": (128, 92, 190),
    "bg": (248, 250, 252),
}


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    if FONT_PATH is None:
        return ImageFont.load_default()
    return ImageFont.truetype(str(FONT_PATH), size=size, index=0)


def text_width(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.ImageFont) -> int:
    box = draw.textbbox((0, 0), text, font=fnt)
    return box[2] - box[0]


def draw_centered(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int],
    text: str,
    fnt: ImageFont.ImageFont,
    fill: tuple[int, int, int],
) -> None:
    x, y = xy
    w = text_width(draw, text, fnt)
    draw.text((x - w // 2, y), text, font=fnt, fill=fill)


def wrap_text(
    draw: ImageDraw.ImageDraw,
    text: str,
    fnt: ImageFont.ImageFont,
    max_width: int,
) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = word if not current else f"{current} {word}"
        if text_width(draw, candidate, fnt) <= max_width:
            current = candidate
            continue
        if current:
            lines.append(current)
        current = word
    if current:
        lines.append(current)
    return lines


def save_bar_chart(
    path: Path,
    title: str,
    labels: Sequence[str],
    values: Sequence[float],
    value_suffix: str = "",
    width: int = 1500,
    height: int = 900,
    color: tuple[int, int, int] = COLORS["blue"],
    y_label: str = "",
) -> None:
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_font = font(42)
    label_font = font(25)
    small_font = font(22)
    draw.text((70, 45), title, font=title_font, fill=COLORS["ink"])
    if y_label:
        draw.text((72, 100), y_label, font=small_font, fill=COLORS["muted"])

    left, top, right, bottom = 120, 165, width - 70, height - 150
    max_value = max(values) if values else 1
    max_value = max(max_value, 1)
    steps = 5
    for i in range(steps + 1):
        y = bottom - (bottom - top) * i / steps
        draw.line((left, y, right, y), fill=COLORS["grid"], width=2)
        tick = max_value * i / steps
        draw.text((22, y - 12), f"{tick:.2f}", font=small_font, fill=COLORS["muted"])

    gap = 28
    bar_w = max(24, int((right - left - gap * (len(labels) + 1)) / max(len(labels), 1)))
    x = left + gap
    for label, value in zip(labels, values):
        bar_h = int((value / max_value) * (bottom - top))
        y0 = bottom - bar_h
        draw.rounded_rectangle((x, y0, x + bar_w, bottom), radius=8, fill=color)
        value_text = f"{value:.3f}{value_suffix}" if value < 10 else f"{value:,.0f}{value_suffix}"
        draw_centered(draw, (x + bar_w // 2, y0 - 34), value_text, small_font, COLORS["ink"])
        label_lines = wrap_text(draw, label, label_font, bar_w + 50)
        for idx, line in enumerate(label_lines[:2]):
            draw_centered(
                draw,
                (x + bar_w // 2, bottom + 18 + idx * 30),
                line,
                label_font,
                COLORS["ink"],
            )
        x += bar_w + gap

    draw.line((left, bottom, right, bottom), fill=COLORS["ink"], width=3)
    img.save(path)


def save_horizontal_bar_chart(
    path: Path,
    title: str,
    labels: Sequence[str],
    values: Sequence[float],
    width: int = 1500,
    height: int = 950,
    color: tuple[int, int, int] = COLORS["green"],
    value_format: str = "{:.3f}",
) -> None:
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_font = font(42)
    label_font = font(25)
    small_font = font(22)
    draw.text((70, 45), title, font=title_font, fill=COLORS["ink"])

    left, top, right, bottom = 390, 140, width - 110, height - 60
    max_value = max(values) if values else 1
    max_value = max(max_value, 1e-9)
    row_h = (bottom - top) / max(len(labels), 1)
    for idx, (label, value) in enumerate(zip(labels, values)):
        y = top + idx * row_h + row_h * 0.18
        bar_h = row_h * 0.55
        draw.text((70, int(y + 5)), label, font=label_font, fill=COLORS["ink"])
        draw.rounded_rectangle((left, y, right, y + bar_h), radius=7, fill=(235, 240, 246))
        x1 = left + int((value / max_value) * (right - left))
        draw.rounded_rectangle((left, y, x1, y + bar_h), radius=7, fill=color)
        draw.text((x1 + 14, int(y + 5)), value_format.format(value), font=small_font, fill=COLORS["ink"])
    img.save(path)


def save_histogram(path: Path, title: str, values: pd.Series) -> None:
    clean = values.dropna().clip(lower=0)
    log_values = clean.map(lambda v: math.log10(v + 1))
    bins = [0, 4, 5, 5.5, 6, 6.3, 6.5, 6.7, 7, 7.3, 7.5]
    counts = []
    labels = []
    for start, end in zip(bins, bins[1:]):
        counts.append(int(((log_values >= start) & (log_values < end)).sum()))
        labels.append(f"{int(10 ** start):,}\n~{int(10 ** end):,}")
    save_bar_chart(
        path,
        title,
        labels,
        counts,
        width=1700,
        height=950,
        color=COLORS["orange"],
        y_label="영화 수, x축은 누적 관객 수 구간",
    )


def save_workflow_diagram(path: Path) -> None:
    width, height = 1700, 850
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_font = font(42)
    box_font = font(25)
    small_font = font(21)
    draw.text((70, 45), "분석 파이프라인", font=title_font, fill=COLORS["ink"])
    boxes = [
        ("KOBIS/TMDB 수집", "관객 수, 개봉일, 장르, 런타임, 포스터 URL"),
        ("제목 매핑", "KOBIS 제목과 TMDB 제목 기준 병합"),
        ("특성 추출", "메타데이터, 제목, 포스터 색상/복잡도"),
        ("모델 학습", "Logistic, Tree, RF, MLP, 회귀 모델"),
        ("평가/해석", "F1, ROC-AUC, RMSE, 변수 중요도"),
    ]
    x, y = 70, 210
    box_w, box_h, gap = 280, 235, 35
    for idx, (heading, body) in enumerate(boxes):
        fill = [COLORS["blue"], COLORS["green"], COLORS["orange"], COLORS["purple"], COLORS["red"]][idx]
        draw.rounded_rectangle((x, y, x + box_w, y + box_h), radius=20, fill=(248, 250, 252), outline=fill, width=5)
        draw.text((x + 28, y + 28), heading, font=box_font, fill=fill)
        for line_idx, line in enumerate(wrap_text(draw, body, small_font, box_w - 56)[:4]):
            draw.text((x + 28, y + 88 + line_idx * 32), line, font=small_font, fill=COLORS["ink"])
        if idx < len(boxes) - 1:
            arrow_x = x + box_w + 5
            mid_y = y + box_h // 2
            draw.line((arrow_x, mid_y, arrow_x + gap - 10, mid_y), fill=COLORS["ink"], width=5)
            draw.polygon(
                [(arrow_x + gap - 10, mid_y - 12), (arrow_x + gap + 10, mid_y), (arrow_x + gap - 10, mid_y + 12)],
                fill=COLORS["ink"],
            )
        x += box_w + gap
    img.save(path)


def poster_path(title: str) -> Path | None:
    direct = ROOT / "data" / "posters" / f"{title}.jpg"
    if direct.exists():
        return direct
    matches = sorted((ROOT / "data" / "posters").glob(f"{title}.*"))
    return matches[0] if matches else None


def save_poster_collage(path: Path, titles: Sequence[str]) -> None:
    thumb_w, thumb_h = 210, 315
    gap = 24
    width = 2 * gap + len(titles) * thumb_w + (len(titles) - 1) * gap
    height = 470
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_font = font(34)
    label_font = font(20)
    draw.text((gap, 25), "포스터 이미지 예시", font=title_font, fill=COLORS["ink"])
    x = gap
    for title in titles:
        p = poster_path(title)
        if p and p.exists():
            poster = Image.open(p).convert("RGB")
            poster.thumbnail((thumb_w, thumb_h), Image.Resampling.LANCZOS)
            canvas = Image.new("RGB", (thumb_w, thumb_h), (238, 242, 247))
            px = (thumb_w - poster.width) // 2
            py = (thumb_h - poster.height) // 2
            canvas.paste(poster, (px, py))
            img.paste(canvas, (x, 90))
        else:
            draw.rounded_rectangle((x, 90, x + thumb_w, 90 + thumb_h), radius=10, fill=(238, 242, 247))
            draw_centered(draw, (x + thumb_w // 2, 225), "이미지 없음", label_font, COLORS["muted"])
        for idx, line in enumerate(wrap_text(draw, title, label_font, thumb_w)[:2]):
            draw_centered(draw, (x + thumb_w // 2, 420 + idx * 24), line, label_font, COLORS["ink"])
        x += thumb_w + gap
    img.save(path)


def create_charts(frame: pd.DataFrame) -> dict[str, Path]:
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)

    paths = {
        "workflow": FIGURE_DIR / "workflow.png",
        "poster_collage": FIGURE_DIR / "poster_collage.png",
        "audience_hist": FIGURE_DIR / "audience_distribution.png",
        "success_ratio": FIGURE_DIR / "success_ratio.png",
        "genre_avg": FIGURE_DIR / "genre_average_audience.png",
        "classification": FIGURE_DIR / "classification_model_comparison.png",
        "regression": FIGURE_DIR / "regression_model_comparison.png",
        "importance": FIGURE_DIR / "feature_importance.png",
        "poster_success": FIGURE_DIR / "poster_feature_success.png",
    }
    save_workflow_diagram(paths["workflow"])
    save_poster_collage(paths["poster_collage"], ["명량", "극한직업", "베테랑", "괴물", "도둑들", "기생충"])
    save_histogram(paths["audience_hist"], "누적 관객 수 분포", frame["audience_count"])

    success_labels = ["100만 이상", "300만 이상", "500만 이상"]
    success_values = [
        frame["is_success_1000000"].mean() * 100,
        frame["is_success_3000000"].mean() * 100,
        frame["is_success_5000000"].mean() * 100,
    ]
    save_bar_chart(
        paths["success_ratio"],
        "흥행 기준별 성공 영화 비율",
        success_labels,
        success_values,
        value_suffix="%",
        color=COLORS["green"],
        y_label="전체 654편 중 기준 관객 수 이상 비율",
    )

    genre_rows = []
    for _, row in frame[["tmdb_genres", "audience_count"]].dropna().iterrows():
        for genre in str(row["tmdb_genres"]).split("|"):
            if genre:
                genre_rows.append({"genre": genre, "audience": row["audience_count"]})
    genre_frame = pd.DataFrame(genre_rows)
    genre_stats = (
        genre_frame.groupby("genre")
        .agg(avg=("audience", "mean"), count=("audience", "size"))
        .query("count >= 20")
        .sort_values("avg", ascending=False)
        .head(10)
    )
    save_horizontal_bar_chart(
        paths["genre_avg"],
        "장르별 평균 관객 수 Top 10",
        genre_stats.index.tolist(),
        (genre_stats["avg"] / 1_000_000).tolist(),
        color=COLORS["orange"],
        value_format="{:.2f}백만",
    )

    cls = pd.read_csv(CLASSIFICATION_METRICS, encoding="utf-8-sig")
    cls_rf = cls[cls["model"].eq("random_forest")].copy()
    labels = [experiment_label(v) for v in cls_rf["experiment"]]
    save_bar_chart(
        paths["classification"],
        "Random Forest 입력 특성 조합별 ROC-AUC",
        labels,
        cls_rf["roc_auc"].tolist(),
        color=COLORS["blue"],
        y_label="300만 이상 흥행 여부 예측",
    )

    reg = pd.read_csv(REGRESSION_METRICS, encoding="utf-8-sig")
    reg_best = reg[reg["model"].isin(["linear_regression", "random_forest_regressor"])].copy()
    reg_best["label"] = reg_best["experiment"].map(experiment_label) + "\n" + reg_best["model"].map(model_label)
    reg_best = reg_best.sort_values("rmse").head(6)
    save_bar_chart(
        paths["regression"],
        "관객 수 회귀 모델 RMSE 비교",
        reg_best["label"].tolist(),
        (reg_best["rmse"] / 1_000_000).tolist(),
        color=COLORS["purple"],
        y_label="백만 명 단위, 낮을수록 좋음",
    )

    importance = pd.read_csv(IMPORTANCE_FILE, encoding="utf-8-sig").head(12)
    save_horizontal_bar_chart(
        paths["importance"],
        "전체 특성 Random Forest 변수 중요도 Top 12",
        [feature_label(v) for v in importance["feature"]],
        importance["importance"].tolist(),
        color=COLORS["red"],
        value_format="{:.3f}",
    )

    poster_metrics = frame.groupby("is_success_3000000").agg(
        brightness=("poster_brightness_mean", "mean"),
        saturation=("poster_saturation_mean", "mean"),
        contrast=("poster_contrast", "mean"),
        edge=("poster_edge_density", "mean"),
    )
    labels = ["비흥행\n밝기", "흥행\n밝기", "비흥행\n채도", "흥행\n채도", "비흥행\n대비", "흥행\n대비"]
    values = [
        poster_metrics.loc[False, "brightness"],
        poster_metrics.loc[True, "brightness"],
        poster_metrics.loc[False, "saturation"],
        poster_metrics.loc[True, "saturation"],
        poster_metrics.loc[False, "contrast"],
        poster_metrics.loc[True, "contrast"],
    ]
    save_bar_chart(
        paths["poster_success"],
        "흥행 여부별 포스터 색상 특성 평균",
        labels,
        values,
        color=COLORS["green"],
        y_label="0~1 정규화 값",
    )
    paths.update(copy_visualization_upgrade_charts())
    paths.update(create_notebook_analysis_charts())
    return paths


def copy_visualization_upgrade_charts() -> dict[str, Path]:
    mapping = {
        "rank_distribution": "rank_distribution.png",
        "movie_feature_pca_map": "movie_feature_pca_map.png",
        "actual_audience_vs_predicted_rank": "actual_audience_vs_predicted_rank.png",
        "rank_confusion_heatmap": "rank_confusion_heatmap.png",
        "genre_rank_heatmap": "genre_rank_heatmap.png",
        "month_rank_heatmap": "month_rank_heatmap.png",
        "poster_brightness_saturation_scatter": "poster_brightness_saturation_scatter.png",
        "misclassified_movies_lollipop": "misclassified_movies_lollipop.png",
    }
    copied: dict[str, Path] = {}
    for key, filename in mapping.items():
        source = VISUALIZATION_UPGRADE_DIR / filename
        target = FIGURE_DIR / filename
        if source.exists():
            shutil.copy2(source, target)
            copied[key] = target
    return copied


def configure_matplotlib():
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.font_manager as fm
    import matplotlib.pyplot as plt

    preferred_fonts = ["AppleGothic", "Apple SD Gothic Neo", "NanumGothic", "Malgun Gothic", "DejaVu Sans"]
    available_fonts = {font.name for font in fm.fontManager.ttflist}
    for family in preferred_fonts:
        if family in available_fonts:
            plt.rcParams["font.family"] = family
            break
    plt.rcParams["axes.unicode_minus"] = False
    plt.rcParams["figure.dpi"] = 130
    plt.rcParams["savefig.dpi"] = 180
    return plt


def split_names(value: object) -> list[str]:
    if pd.isna(value):
        return []
    text = str(value).strip()
    if not text or text.casefold() in {"nan", "none", "null"}:
        return []
    return [clean_name(part) for part in text.split("|") if clean_name(part)]


def clean_name(value: object) -> str:
    return re.sub(r"\([^)]*\)$", "", str(value).strip()).strip()


def actor_movie_frame() -> pd.DataFrame:
    raw = pd.read_csv(RAW_DATA_FILE, encoding="utf-8-sig")
    title_col = "match_title" if "match_title" in raw.columns else raw.columns[0]
    result = pd.DataFrame(
        {
            "movie_title": raw[title_col],
            "audience_count": pd.to_numeric(raw["audience_count"], errors="coerce"),
            "actors": raw["tmdb_cast"].map(split_names),
        }
    )
    result = result.dropna(subset=["audience_count"])
    result = result[(result["audience_count"] > 0) & result["actors"].map(bool)].copy()
    result["audience_million"] = result["audience_count"] / 1_000_000
    return result.reset_index(drop=True)


def build_actor_matrix(movies: pd.DataFrame, min_movie_count: int = 3) -> tuple[pd.DataFrame, pd.DataFrame]:
    actor_counts = pd.Series(
        [actor for actors in movies["actors"] for actor in actors],
        dtype="object",
    ).value_counts()
    selected_actors = actor_counts[actor_counts >= min_movie_count].index.tolist()
    matrix = pd.DataFrame(0, index=movies.index, columns=selected_actors, dtype=int)
    actor_set = set(selected_actors)
    for idx, actors in movies["actors"].items():
        for actor in actors:
            if actor in actor_set:
                matrix.at[idx, actor] = 1
    matrix = matrix.loc[:, matrix.sum().sort_values(ascending=False).index]

    rows = []
    for actor in matrix.columns:
        mask = matrix[actor].eq(1)
        audiences = movies.loc[mask, "audience_count"]
        rows.append(
            {
                "actor": actor,
                "n_movies": int(mask.sum()),
                "mean_audience": float(audiences.mean()),
                "median_audience": float(audiences.median()),
                "hit_rate_3m": float((audiences >= 3_000_000).mean()),
                "hit_rate_5m": float((audiences >= 5_000_000).mean()),
            }
        )
    stats = pd.DataFrame(rows).set_index("actor")
    return matrix, stats


def create_actor_lasso_chart(path: Path) -> dict[str, object]:
    import numpy as np
    from sklearn.dummy import DummyRegressor
    from sklearn.linear_model import LassoCV
    from sklearn.metrics import mean_absolute_error, mean_squared_error, r2_score
    from sklearn.model_selection import train_test_split

    plt = configure_matplotlib()
    movies = actor_movie_frame()
    matrix, actor_stats = build_actor_matrix(movies)
    y = np.log1p(movies["audience_count"])
    x_train, x_test, y_train, y_test = train_test_split(matrix, y, test_size=0.2, random_state=42)

    baseline = DummyRegressor(strategy="mean")
    baseline.fit(x_train, y_train)
    baseline_pred = baseline.predict(x_test)

    lasso = LassoCV(alphas=np.logspace(-4, 0, 80), cv=5, max_iter=20_000, random_state=42)
    lasso.fit(x_train, y_train)
    lasso_pred = lasso.predict(x_test)

    def metrics(y_true_log, y_pred_log) -> dict[str, float]:
        actual = np.expm1(y_true_log)
        predicted = np.expm1(y_pred_log).clip(min=0)
        return {
            "mae": float(mean_absolute_error(actual, predicted)),
            "rmse": float(np.sqrt(mean_squared_error(actual, predicted))),
            "r2": float(r2_score(actual, predicted)),
        }

    metric_rows = pd.DataFrame(
        [
            {"model": "평균 예측 기준선", **metrics(y_test, baseline_pred)},
            {"model": "LassoCV 배우 원-핫", **metrics(y_test, lasso_pred)},
        ]
    )
    coef_summary = pd.DataFrame(
        {
            "actor": matrix.columns,
            "lasso_coef_log": lasso.coef_,
        }
    ).join(actor_stats, on="actor")
    baseline_log_audience = float(y_train.mean())
    coef_summary["estimated_delta_audience"] = (
        np.expm1(baseline_log_audience + coef_summary["lasso_coef_log"])
        - np.expm1(baseline_log_audience)
    )
    coef_summary["estimated_delta_million"] = coef_summary["estimated_delta_audience"] / 1_000_000
    coef_summary["mean_audience_million"] = coef_summary["mean_audience"] / 1_000_000
    coef_summary["median_audience_million"] = coef_summary["median_audience"] / 1_000_000
    selected = coef_summary[coef_summary["lasso_coef_log"].abs() > 1e-8].copy()
    if selected.empty:
        selected = coef_summary.reindex(coef_summary["lasso_coef_log"].abs().sort_values(ascending=False).index).head(20)

    plot_data = selected.sort_values("lasso_coef_log", ascending=False).head(15)
    plot_data = plot_data.sort_values("estimated_delta_million")
    fig, ax = plt.subplots(figsize=(10, 6))
    colors = np.where(plot_data["estimated_delta_million"] >= 0, "#2A9D8F", "#E76F51")
    ax.barh(plot_data["actor"], plot_data["estimated_delta_million"], color=colors)
    ax.axvline(0, color="#333333", linewidth=1)
    ax.set_title("Lasso가 선택한 배우별 관객 수 연관 신호")
    ax.set_xlabel("평균 로그 관객 대비 추정 차이 (백만 명)")
    ax.set_ylabel("")
    ax.grid(axis="x", alpha=0.25)
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return {
        "metrics": metric_rows,
        "top_positive": selected.sort_values("lasso_coef_log", ascending=False).head(8),
        "n_movies": len(movies),
        "n_actors": len(matrix.columns),
    }


def create_actor_hit_probability_chart(path: Path) -> dict[str, object]:
    import numpy as np
    from sklearn.dummy import DummyClassifier
    from sklearn.linear_model import LogisticRegressionCV
    from sklearn.metrics import accuracy_score, f1_score, precision_score, recall_score, roc_auc_score
    from sklearn.model_selection import train_test_split

    plt = configure_matplotlib()
    movies = actor_movie_frame()
    matrix, actor_stats = build_actor_matrix(movies)
    y = (movies["audience_count"] >= 5_000_000).astype(int)
    x_train, x_test, y_train, y_test = train_test_split(
        matrix,
        y,
        test_size=0.2,
        random_state=42,
        stratify=y,
    )

    baseline = DummyClassifier(strategy="most_frequent")
    baseline.fit(x_train, y_train)
    cv = min(5, int(y_train.value_counts().min()))
    logit = LogisticRegressionCV(
        Cs=np.logspace(-2, 2, 12),
        cv=cv,
        scoring="roc_auc" if cv >= 2 else "accuracy",
        class_weight="balanced",
        max_iter=5_000,
        random_state=42,
    )
    logit.fit(x_train, y_train)

    def metrics(name: str, estimator) -> dict[str, float | str]:
        pred = estimator.predict(x_test)
        prob = estimator.predict_proba(x_test)[:, 1] if hasattr(estimator, "predict_proba") else None
        return {
            "model": name,
            "accuracy": float(accuracy_score(y_test, pred)),
            "precision": float(precision_score(y_test, pred, zero_division=0)),
            "recall": float(recall_score(y_test, pred, zero_division=0)),
            "f1": float(f1_score(y_test, pred, zero_division=0)),
            "roc_auc": float(roc_auc_score(y_test, prob)) if prob is not None and y_test.nunique() == 2 else math.nan,
        }

    single_actor_matrix = pd.DataFrame(np.eye(len(matrix.columns)), columns=matrix.columns)
    single_actor_prob = logit.predict_proba(single_actor_matrix)[:, 1]
    actor_signal = pd.DataFrame(
        {
            "actor": matrix.columns,
            "logit_coef": logit.coef_[0],
            "single_actor_hit_prob": single_actor_prob,
        }
    ).join(actor_stats, on="actor")
    actor_signal["single_actor_hit_prob_pct"] = actor_signal["single_actor_hit_prob"] * 100
    actor_signal["mean_audience_million"] = actor_signal["mean_audience"] / 1_000_000
    actor_signal["median_audience_million"] = actor_signal["median_audience"] / 1_000_000
    actor_signal["profile_score"] = (
        actor_signal["n_movies"].rank(pct=True) * 0.55
        + actor_signal["mean_audience_million"].rank(pct=True) * 0.45
    )
    actor_signal["profile_group"] = np.select(
        [
            actor_signal["profile_score"] >= actor_signal["profile_score"].quantile(0.67),
            actor_signal["profile_score"] <= actor_signal["profile_score"].quantile(0.40),
        ],
        ["상위 인지도", "낮은 인지도"],
        default="중간 인지도",
    )

    top_probability = actor_signal.sort_values("single_actor_hit_prob_pct", ascending=False).head(8)
    lower_profile_probability = (
        actor_signal[
            actor_signal["profile_group"].ne("상위 인지도")
            & actor_signal["n_movies"].le(8)
            & ~actor_signal["actor"].isin(top_probability["actor"])
        ]
        .sort_values("single_actor_hit_prob_pct", ascending=False)
        .head(8)
    )

    plot_data = pd.concat(
        [
            top_probability.assign(display_group="상위 확률"),
            lower_profile_probability.assign(display_group="중저인지도 후보"),
        ],
        ignore_index=True,
    ).drop_duplicates(subset=["actor"])
    plot_data["actor_label"] = plot_data["actor"] + " · " + plot_data["display_group"]
    plot_data = plot_data.sort_values(["display_group", "single_actor_hit_prob_pct"], ascending=[True, True])
    fig, ax = plt.subplots(figsize=(10, 6))
    colors = plot_data["display_group"].map({"상위 확률": "#457B9D", "중저인지도 후보": "#E07A5F"}).fillna("#457B9D")
    ax.barh(plot_data["actor_label"], plot_data["single_actor_hit_prob_pct"], color=colors)
    ax.set_xlim(0, 100)
    ax.set_title("배우 신호만 넣었을 때 500만 돌파 확률 후보")
    ax.set_xlabel("예측 확률 (%)")
    ax.set_ylabel("")
    ax.grid(axis="x", alpha=0.25)
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return {
        "metrics": pd.DataFrame([metrics("최빈값 기준선", baseline), metrics("LogisticCV 배우 원-핫", logit)]),
        "top_probability": top_probability,
        "lower_profile_probability": lower_profile_probability,
        "class_counts": y.value_counts().to_dict(),
    }


def create_budget_cluster_chart(path: Path) -> dict[str, object]:
    import numpy as np
    from sklearn.cluster import KMeans
    from sklearn.metrics import silhouette_score
    from sklearn.preprocessing import StandardScaler

    plt = configure_matplotlib()
    raw = pd.read_csv(DATA_FILE, encoding="utf-8-sig")
    movies = pd.DataFrame(
        {
            "movie_title": raw["match_title"],
            "budget": pd.to_numeric(raw["meta_tmdb_budget"], errors="coerce"),
            "audience_count": pd.to_numeric(raw["audience_count"], errors="coerce"),
        }
    )
    movies = movies.dropna(subset=["budget", "audience_count"])
    movies = movies[(movies["budget"] > 0) & (movies["audience_count"] > 0)].copy()
    movies["budget_million"] = movies["budget"] / 1_000_000
    movies["audience_million"] = movies["audience_count"] / 1_000_000
    movies["log_budget"] = np.log1p(movies["budget"])
    movies["log_audience"] = np.log1p(movies["audience_count"])
    movies["audience_per_budget_million"] = movies["audience_million"] / movies["budget_million"]

    features = movies[["log_budget", "log_audience"]]
    scaled = StandardScaler().fit_transform(features)
    max_k = min(7, len(movies) - 1)
    silhouette_rows = []
    for k in range(2, max_k + 1):
        labels = KMeans(n_clusters=k, random_state=42, n_init=20).fit_predict(scaled)
        silhouette_rows.append({"k": k, "silhouette": float(silhouette_score(scaled, labels))})
    silhouette_table = pd.DataFrame(silhouette_rows)
    selected_k = 4 if len(movies) >= 4 else int(silhouette_table.sort_values("silhouette", ascending=False).iloc[0]["k"])
    movies["cluster"] = KMeans(n_clusters=selected_k, random_state=42, n_init=20).fit_predict(scaled)

    summary = movies.groupby("cluster").agg(
        n_movies=("movie_title", "count"),
        median_budget_million=("budget_million", "median"),
        median_audience_million=("audience_million", "median"),
        mean_audience_million=("audience_million", "mean"),
        hit_rate_3m=("audience_count", lambda values: (values >= 3_000_000).mean()),
        median_audience_per_budget_million=("audience_per_budget_million", "median"),
    )
    overall_budget = movies["budget_million"].median()
    overall_audience = movies["audience_million"].median()

    def cluster_label(row: pd.Series) -> str:
        budget_high = row["median_budget_million"] >= overall_budget
        audience_high = row["median_audience_million"] >= overall_audience
        if budget_high and audience_high:
            return "대작 흥행형"
        if budget_high and not audience_high:
            return "고비용 저효율형"
        if not budget_high and audience_high:
            return "저비용 고효율형"
        return "소형 저관객형"

    summary["cluster_label"] = summary.apply(cluster_label, axis=1)
    if summary["cluster_label"].duplicated().any():
        summary["cluster_label"] = [
            f"군집 {cluster_id + 1}: {label}"
            for cluster_id, label in zip(summary.index, summary["cluster_label"])
        ]
    movies["cluster_label"] = movies["cluster"].map(summary["cluster_label"])

    fig, ax = plt.subplots(figsize=(11, 7))
    for label, part in movies.groupby("cluster_label"):
        ax.scatter(
            part["budget_million"],
            part["audience_million"],
            s=70,
            alpha=0.75,
            label=f"{label} ({len(part)})",
        )
    ax.axhline(3, color="#555555", linestyle="--", linewidth=1, label="300만 기준")
    ax.set_xscale("log")
    ax.set_yscale("log")
    ax.set_title("제작비-관객 수 기준 영화 군집")
    ax.set_xlabel("TMDB 제작비 (백만 단위, 로그축)")
    ax.set_ylabel("누적 관객 수 (백만 명, 로그축)")
    ax.grid(True, which="both", alpha=0.2)
    ax.legend(loc="best")
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)

    efficiency = movies.sort_values(
        ["audience_per_budget_million", "audience_count"],
        ascending=False,
    ).head(8)
    return {
        "n_movies": len(movies),
        "selected_k": selected_k,
        "silhouette": silhouette_table,
        "summary": summary.sort_values("median_audience_million", ascending=False),
        "efficiency": efficiency,
    }


def create_notebook_analysis_charts() -> dict[str, Path]:
    paths = {
        "actor_lasso_signal": FIGURE_DIR / "actor_lasso_signal.png",
        "actor_hit_probability": FIGURE_DIR / "actor_hit_probability.png",
        "budget_audience_cluster": FIGURE_DIR / "budget_audience_cluster.png",
    }
    create_actor_lasso_chart(paths["actor_lasso_signal"])
    create_actor_hit_probability_chart(paths["actor_hit_probability"])
    create_budget_cluster_chart(paths["budget_audience_cluster"])
    return paths


def experiment_label(value: str) -> str:
    return {
        "metadata": "메타",
        "metadata_title": "메타+제목",
        "metadata_poster": "메타+포스터",
        "all": "전체",
    }.get(value, value)


def model_label(value: str) -> str:
    return {
        "logistic_regression": "Logistic",
        "decision_tree": "Tree",
        "random_forest": "RF",
        "dnn_mlp": "MLP",
        "linear_regression": "Linear",
        "ridge": "Ridge",
        "lasso": "Lasso",
        "random_forest_regressor": "RF Regr.",
    }.get(value, value)


def feature_label(value: str) -> str:
    label_map = {
        "meta_tmdb_runtime": "런타임",
        "meta_cast_mean_frequency": "출연진 평균 빈도",
        "meta_cast_max_frequency": "출연진 최대 빈도",
        "meta_log_tmdb_budget": "제작비 로그",
        "meta_tmdb_budget": "제작비",
        "meta_open_year": "개봉 연도",
        "meta_director_mean_frequency": "감독 평균 빈도",
        "meta_director_max_frequency": "감독 최대 빈도",
        "meta_tmdb_release_year": "TMDB 공개 연도",
        "poster_red_mean": "포스터 R 평균",
        "poster_brightness_mean": "포스터 밝기",
        "poster_green_mean": "포스터 G 평균",
        "poster_blue_mean": "포스터 B 평균",
        "poster_saturation_mean": "포스터 채도",
        "poster_contrast": "포스터 대비",
        "poster_edge_density": "포스터 에지 밀도",
    }
    return label_map.get(value, re.sub(r"^(meta|poster|title)_", "", value).replace("_", " "))


def table_data(frame: pd.DataFrame) -> dict[str, object]:
    cls = pd.read_csv(CLASSIFICATION_METRICS, encoding="utf-8-sig")
    reg = pd.read_csv(REGRESSION_METRICS, encoding="utf-8-sig")
    rank = pd.read_csv(RANK_CLASSIFICATION_METRICS, encoding="utf-8-sig")
    rank_predictions = pd.read_csv(RANK_CLASSIFICATION_PREDICTIONS, encoding="utf-8-sig")
    actor_lasso = create_actor_lasso_chart(FIGURE_DIR / "actor_lasso_signal.png")
    actor_hit = create_actor_hit_probability_chart(FIGURE_DIR / "actor_hit_probability.png")
    budget_cluster = create_budget_cluster_chart(FIGURE_DIR / "budget_audience_cluster.png")
    best_rank = rank.sort_values(["f1_macro", "accuracy"], ascending=False).iloc[0]
    best_rank_predictions = rank_predictions[
        (rank_predictions["experiment"] == best_rank["experiment"])
        & (rank_predictions["model"] == best_rank["model"])
    ].copy()
    return {
        "rows": len(frame),
        "year_min": int(frame["meta_open_year"].min()),
        "year_max": int(frame["meta_open_year"].max()),
        "poster_available": int(frame["poster_available"].sum()),
        "success_1m": int(frame["is_success_1000000"].sum()),
        "success_3m": int(frame["is_success_3000000"].sum()),
        "success_5m": int(frame["is_success_5000000"].sum()),
        "audience_mean": int(frame["audience_count"].mean()),
        "audience_median": int(frame["audience_count"].median()),
        "audience_max": int(frame["audience_count"].max()),
        "classification": cls,
        "regression": reg,
        "rank": rank,
        "rank_predictions": best_rank_predictions,
        "best_cls_roc": cls.sort_values("roc_auc", ascending=False).iloc[0],
        "best_cls_f1": cls.sort_values("f1", ascending=False).iloc[0],
        "best_cls_accuracy": cls.sort_values("accuracy", ascending=False).iloc[0],
        "best_reg_r2": reg.sort_values("r2", ascending=False).iloc[0],
        "best_reg_rmse": reg.sort_values("rmse", ascending=True).iloc[0],
        "best_rank": best_rank,
        "rank_counts": frame["success_rank"].value_counts().reindex(RANK_ORDER).fillna(0).astype(int),
        "actor_lasso": actor_lasso,
        "actor_hit": actor_hit,
        "budget_cluster": budget_cluster,
    }


def set_cell_text(cell, text: object, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(str(text))
    set_run_font(run, size=8.5, bold=bold, color=RGBColor(38, 45, 56))


def qn(tag: str):
    from docx.oxml.ns import qn as docx_qn

    return docx_qn(tag)


def set_run_font(
    run,
    size: float = 10.5,
    bold: bool | None = None,
    color: RGBColor | None = None,
    name: str = "Apple SD Gothic Neo",
) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.first_child_found_in("w:shd")
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_width(cell, width_inches: float) -> None:
    width = Inches(width_inches)
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.twips)))
    tc_w.set(qn("w:type"), "dxa")


def add_table(
    doc: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[object]],
    widths: Sequence[float] | None = None,
) -> None:
    if widths is None:
        widths = [6.5 / max(len(headers), 1)] * len(headers)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, header, bold=True)
        set_cell_width(cell, widths[idx])
        set_cell_margins(cell)
        set_cell_shading(cell, "F2F4F7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            set_cell_text(cells[idx], value)
            set_cell_width(cells[idx], widths[idx])
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    paragraph = doc.add_paragraph(style=style)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.10
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, color=RGBColor(38, 45, 56))


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.08)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.10
        run = paragraph.add_run(item)
        set_run_font(run, size=10, color=RGBColor(38, 45, 56))


def add_picture(doc: Document, path: Path, width: float = 6.2, caption: str | None = None) -> None:
    if not path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        set_run_font(run, size=8.8, color=RGBColor(102, 112, 133))
        cap.paragraph_format.space_after = Pt(8)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.keep_with_next = True
    heading.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    heading.paragraph_format.space_after = Pt(8 if level == 1 else 5)
    for run in heading.runs:
        set_run_font(
            run,
            size=16 if level == 1 else 13 if level == 2 else 12,
            bold=True,
            color=RGBColor(46, 116, 181) if level <= 2 else RGBColor(31, 77, 120),
        )


def classification_rows(cls: pd.DataFrame) -> list[list[str]]:
    rows = []
    for _, row in cls.iterrows():
        rows.append(
            [
                experiment_label(row["experiment"]),
                model_label(row["model"]),
                f"{row['accuracy']:.3f}",
                f"{row['precision']:.3f}",
                f"{row['recall']:.3f}",
                f"{row['f1']:.3f}",
                f"{row['roc_auc']:.3f}",
            ]
        )
    return rows


def best_rows(frame: pd.DataFrame, metric: str, ascending: bool = False, n: int = 5) -> list[list[str]]:
    rows = []
    for _, row in frame.sort_values(metric, ascending=ascending).head(n).iterrows():
        strategy = row.get("strategy", "")
        rows.append(
            [
                strategy if isinstance(strategy, str) else "",
                experiment_label(row["experiment"]),
                model_label(row["model"]),
                f"{row[metric]:.3f}" if abs(row[metric]) < 100 else f"{row[metric]:,.0f}",
            ]
        )
    return rows


def build_docx(frame: pd.DataFrame, charts: dict[str, Path], stats: dict[str, object]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = doc.styles
    styles["Normal"].font.name = "Apple SD Gothic Neo"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(90)
    run = title.add_run("국내 영화 흥행 여부 예측 및 성공 요인 분석")
    run.bold = True
    run.font.name = "Apple SD Gothic Neo"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(38, 45, 56)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("영화 제목, 포스터, 메타데이터를 중심으로")
    run.font.name = "Apple SD Gothic Neo"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(80, 88, 100)
    add_picture(doc, charts["poster_collage"], width=6.0)
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["교과목명", "데이터마이닝"],
            ["담당 교수명", "[기입]"],
            ["팀명", "[기입]"],
            ["팀원 학번 / 이름", "[기입]"],
            ["제출일", "2026년 6월"],
        ],
    )

    doc.add_page_break()
    add_heading(doc, "목차 및 요약", 1)
    add_paragraph(
        doc,
        "본 보고서는 KOBIS 박스오피스 관객 수와 TMDB 영화 메타데이터를 결합해 국내 영화의 흥행 여부를 예측한 데이터마이닝 프로젝트 결과를 정리한다. "
        "분석 대상은 총 654편이며, 기본 흥행 기준은 누적 관객 수 300만 명 이상으로 설정하였다.",
    )
    add_table(
        doc,
        ["구분", "내용"],
        [
            ["핵심 질문", "메타데이터, 제목 특성, 포스터 특성이 흥행 예측에 도움이 되는가"],
            ["주요 데이터", "KOBIS 관객 수, TMDB 장르·런타임·제작비·출연진·감독·포스터"],
            ["분류 모델", "Logistic Regression, Decision Tree, Random Forest, DNN MLP"],
            ["회귀 모델", "Linear, Ridge, Lasso, Random Forest Regressor"],
            ["주요 결과", "Random Forest가 분류에서 가장 안정적이며, 제목·포스터 특성의 추가 효과는 제한적"],
        ],
    )
    add_bullets(
        doc,
        [
            "프로젝트 개요와 연구 필요성",
            "데이터 수집, 전처리, 특성 설계",
            "모델 설계 및 평가 방법",
            "분류·회귀 실험 결과와 주요 해석",
            "한계점, 향후 개선 방향, 참고문헌, 팀원 기여도",
        ],
    )

    doc.add_page_break()
    add_heading(doc, "1. 프로젝트 개요", 1)
    add_paragraph(
        doc,
        "영화 흥행은 작품성뿐 아니라 장르, 개봉 시기, 출연진 인지도, 제작 규모, 포스터의 시각적 인상 등 여러 요인이 복합적으로 작용한다. "
        "본 프로젝트는 이러한 요인을 정량화해 국내 영화의 흥행 여부를 예측하고, 예측에 기여하는 주요 요인을 분석하는 것을 목표로 한다.",
    )
    add_bullets(
        doc,
        [
            "연구 배경: 영화 시장에서는 개봉 전 흥행 가능성을 예측하려는 수요가 크다.",
            "필요성: 흥행 요인을 정량적으로 비교하면 투자·배급·마케팅 의사결정에 참고할 수 있다.",
            "수행 목표: KOBIS와 TMDB 데이터를 결합하고, 제목·포스터·메타데이터 특성을 비교 실험한다.",
            "최종 산출물: 데이터 전처리 파이프라인, 분류/회귀 모델 성능표, 시각화, 분석 보고서.",
        ],
    )
    add_picture(doc, charts["workflow"], width=6.4, caption="그림 1. 프로젝트 분석 파이프라인")

    doc.add_page_break()
    add_heading(doc, "2. 데이터 설명", 1)
    add_paragraph(
        doc,
        f"분석 데이터는 총 {stats['rows']}편의 국내 영화로 구성된다. 개봉 연도 범위는 {stats['year_min']}년부터 {stats['year_max']}년까지이며, "
        f"포스터 특성은 {stats['poster_available']}편에서 추출되었다. 누적 관객 수 평균은 {stats['audience_mean']:,}명, 중앙값은 {stats['audience_median']:,}명이다.",
    )
    add_table(
        doc,
        ["항목", "값"],
        [
            ["전체 영화 수", f"{stats['rows']:,}편"],
            ["포스터 특성 추출 완료", f"{stats['poster_available']:,}편"],
            ["100만 이상 영화", f"{stats['success_1m']:,}편"],
            ["300만 이상 영화", f"{stats['success_3m']:,}편"],
            ["500만 이상 영화", f"{stats['success_5m']:,}편"],
            ["최대 관객 수", f"{stats['audience_max']:,}명"],
        ],
    )
    add_picture(doc, charts["audience_hist"], width=6.4, caption="그림 2. 누적 관객 수 분포")

    doc.add_page_break()
    add_heading(doc, "3. 전처리 및 특성 설계", 1)
    add_paragraph(
        doc,
        "전처리 단계에서는 관객 수를 분류 타깃과 회귀 타깃으로 변환하고, 수치형·범주형·날짜형·텍스트형·이미지형 변수를 모델 입력으로 바꾸었다. "
        "누적 관객 수는 치우친 분포를 가지므로 회귀 실험에는 log1p 변환을 적용하였다.",
    )
    add_table(
        doc,
        ["특성 그룹", "주요 변수", "처리 방식"],
        [
            ["메타데이터", "런타임, 제작비, 개봉 연도/월, 계절", "결측치 보정, 로그 변환, 계절 더미"],
            ["장르", "액션, 드라마, 코미디 등", "다중 라벨 원-핫 인코딩"],
            ["인물/제작사", "감독, 출연진, 제작사", "등장 빈도 기반 요약 변수"],
            ["제목", "길이, 단어 수, 숫자/영어/특수문자, 키워드", "규칙 기반 텍스트 특성"],
            ["포스터", "밝기, 채도, RGB 평균, 대비, 에지 밀도", "이미지 픽셀 기반 저수준 특성"],
        ],
    )
    add_picture(doc, charts["success_ratio"], width=6.3, caption="그림 3. 흥행 기준별 성공 영화 비율")

    doc.add_page_break()
    add_heading(doc, "4. 관련 이론 및 기술", 1)
    add_paragraph(
        doc,
        "분류 모델은 관객 수 300만 명 이상 여부를 예측하고, 회귀 모델은 로그 변환된 누적 관객 수를 예측한다. "
        "비선형성과 변수 간 상호작용을 확인하기 위해 선형 모델과 트리 기반 모델, 신경망 모델을 함께 비교하였다.",
    )
    add_table(
        doc,
        ["모델", "활용 목적", "해석 포인트"],
        [
            ["Logistic Regression", "기본 분류 기준선", "계수 기반 영향 방향 해석 가능"],
            ["Decision Tree", "규칙 기반 분류", "흥행/비흥행 분기 조건 확인 가능"],
            ["Random Forest", "주요 성능 모델", "비선형 관계와 변수 중요도 확인"],
            ["DNN MLP", "비선형 결합 실험", "복잡한 특성 조합 학습 가능"],
            ["Ridge/Lasso", "규제 회귀", "다중공선성과 변수 선택 효과 확인"],
        ],
    )
    add_picture(doc, charts["genre_avg"], width=6.3, caption="그림 4. 장르별 평균 관객 수")

    doc.add_page_break()
    add_heading(doc, "5. 모델 설계 및 실험 방법", 1)
    add_paragraph(
        doc,
        "입력 특성의 기여도를 확인하기 위해 ablation study를 수행하였다. 동일한 train/test 분할을 사용하고, 메타데이터만 사용한 실험과 제목·포스터 특성을 추가한 실험을 비교하였다.",
    )
    add_table(
        doc,
        ["실험", "입력 특성"],
        [
            ["metadata", "장르, 런타임, 제작비, 개봉 시기, 감독/출연진/제작사 빈도"],
            ["metadata_title", "metadata + 제목 길이, 단어 수, 숫자/영어/특수문자, 키워드"],
            ["metadata_poster", "metadata + 포스터 밝기, 채도, RGB 평균, 대비, 에지 밀도"],
            ["all", "metadata + 제목 특성 + 포스터 특성"],
        ],
    )
    add_bullets(
        doc,
        [
            "학습/평가 분할: 전체 데이터의 80%를 학습, 20%를 테스트로 사용",
            "분류 평가지표: Accuracy, Precision, Recall, F1-score, ROC-AUC, Confusion Matrix",
            "회귀 평가지표: MAE, RMSE, R2, RMSLE",
            "클래스 불균형을 고려해 Accuracy보다 F1과 ROC-AUC도 함께 확인",
        ],
    )

    doc.add_page_break()
    add_heading(doc, "6. 분류 모델 결과", 1)
    add_paragraph(
        doc,
        "300만 명 이상 흥행 여부 예측에서는 Random Forest 계열이 가장 안정적인 성능을 보였다. "
        "메타데이터만 사용한 Random Forest는 Accuracy 0.824, F1 0.489, ROC-AUC 0.818을 기록했고, "
        "메타데이터+제목 Random Forest는 ROC-AUC 0.822로 가장 높았다.",
    )
    cls = stats["classification"]
    add_table(
        doc,
        ["실험", "모델", "Acc", "Prec", "Recall", "F1", "ROC-AUC"],
        classification_rows(cls),
    )
    add_picture(doc, charts["classification"], width=6.2, caption="그림 5. Random Forest 입력 특성 조합별 ROC-AUC")

    doc.add_page_break()
    add_heading(doc, "7. 혼동행렬 및 오분류 해석", 1)
    add_paragraph(
        doc,
        "가장 높은 Accuracy를 보인 metadata Random Forest는 비흥행 영화 판별에는 강했지만, 흥행작 Recall은 0.407로 제한적이었다. "
        "즉, 전체 정확도는 높지만 실제 흥행작 일부를 비흥행으로 예측하는 문제가 남아 있다. 흥행작 탐지가 중요하다면 Recall을 높이는 방향으로 임계값 조정이 필요하다.",
    )
    add_picture(doc, CONFUSION_MATRIX, width=4.8, caption="그림 6. metadata Random Forest 혼동행렬")
    add_table(
        doc,
        ["기준", "상위 모델", "해석"],
        [
            ["Accuracy", "metadata / Random Forest", "비흥행 다수 클래스 판별에 강함"],
            ["F1", "all / Decision Tree", "흥행작 Recall이 높지만 False Positive 증가"],
            ["ROC-AUC", "metadata_title / Random Forest", "확률 순위화 성능이 가장 안정적"],
        ],
    )

    doc.add_page_break()
    add_heading(doc, "8. 회귀 모델 결과", 1)
    reg = stats["regression"]
    add_paragraph(
        doc,
        "누적 관객 수 회귀에서는 전체 특성을 사용한 Linear Regression이 RMSE와 R2 기준으로 가장 우수했다. "
        "다만 R2가 0.211 수준으로 높지 않아 관객 수의 절대 규모를 정확히 예측하기는 어렵다. 이는 흥행 성과가 소수 블록버스터에 크게 치우쳐 있고, 마케팅비나 배급 규모 같은 핵심 외부 변수가 누락되었기 때문이다.",
    )
    add_table(
        doc,
        ["전략", "실험", "모델", "R2"],
        best_rows(reg, "r2", ascending=False, n=5),
    )
    add_picture(doc, charts["regression"], width=6.3, caption="그림 7. 관객 수 회귀 RMSE 비교")

    doc.add_page_break()
    add_heading(doc, "9. 주요 요인 분석", 1)
    add_paragraph(
        doc,
        "전체 특성을 사용한 Random Forest 변수 중요도에서는 런타임, 출연진 빈도, 제작비, 개봉 연도, 감독 빈도, 포스터 색상 특성이 상위권에 나타났다. "
        "이는 흥행이 단일 요인보다 작품 규모, 제작·출연 네트워크, 개봉 시기, 시각적 인상 등이 함께 작용하는 문제임을 보여준다.",
    )
    add_picture(doc, charts["importance"], width=6.3, caption="그림 8. 변수 중요도 Top 12")
    add_picture(doc, charts["poster_success"], width=6.2, caption="그림 9. 흥행 여부별 포스터 색상 특성 평균")

    doc.add_page_break()
    add_heading(doc, "10. 프로젝트 수행 결과 종합", 1)
    rank = stats["rank"]
    add_paragraph(
        doc,
        "최종적으로 본 프로젝트는 데이터 수집부터 특성 추출, 모델링, 평가, 해석까지 영화 흥행 예측의 전체 흐름을 구현하였다. "
        "흥행 여부 분류는 회귀보다 실용성이 높았고, Random Forest가 가장 안정적인 기준 모델로 확인되었다.",
    )
    add_table(
        doc,
        ["전략", "실험", "모델", "F1"],
        best_rows(rank, "f1", ascending=False, n=5),
    )
    add_bullets(
        doc,
        [
            "메타데이터만으로도 일정 수준의 예측력이 확보되었다.",
            "제목 특성은 ROC-AUC를 소폭 높였지만, Accuracy/F1 개선은 제한적이었다.",
            "포스터 특성은 일부 변수 중요도에서 확인되지만 단독으로 큰 성능 향상을 만들지는 못했다.",
            "흥행작 Recall이 낮아 실제 의사결정용 모델로 쓰려면 임계값 조정과 불균형 보정이 필요하다.",
        ],
    )

    doc.add_page_break()
    add_heading(doc, "11. 한계점 및 향후 개선 방향", 1)
    add_paragraph(
        doc,
        "본 프로젝트는 학부 데이터마이닝 프로젝트 수준에서 전체 파이프라인을 완성했다는 점에 의의가 있다. "
        "다만 데이터 수와 변수 범위가 제한적이므로 실제 산업 적용에는 추가 개선이 필요하다.",
    )
    add_table(
        doc,
        ["한계점", "개선 방향"],
        [
            ["데이터 654편으로 표본이 제한됨", "연도별 최신 KOBIS 데이터와 해외 성과 데이터 추가"],
            ["일부 제목 매핑 오류 가능성", "영화 코드·개봉일·감독 정보를 함께 쓰는 정교한 매칭"],
            ["제작비 결측과 마케팅비 부재", "배급사, 스크린 수, 광고비, 예매량, SNS 반응 추가"],
            ["포스터는 저수준 색상 특성 중심", "CLIP/CNN 임베딩 등 고차원 이미지 특성 도입"],
            ["클래스 불균형", "SMOTE, threshold tuning, cost-sensitive learning 적용"],
        ],
    )

    doc.add_page_break()
    add_heading(doc, "12. 참고문헌 및 팀원 기여도", 1)
    add_table(
        doc,
        ["구분", "자료"],
        [
            ["KOBIS", "영화관입장권통합전산망 박스오피스 및 연도별 통계"],
            ["TMDB", "The Movie Database API: 영화 상세 정보, 장르, 포스터, 출연진"],
            ["scikit-learn", "Logistic Regression, Decision Tree, Random Forest, 회귀 모델 구현"],
            ["Pillow", "포스터 이미지 특성 추출 및 보고서 그래프 생성"],
        ],
    )
    add_paragraph(doc, "팀원별 기여도는 실제 팀 구성에 맞춰 아래 표를 수정하면 된다.")
    add_table(
        doc,
        ["팀원", "기여도", "역할"],
        [
            ["[학번/이름]", "[%]", "데이터 수집 및 KOBIS/TMDB 매핑"],
            ["[학번/이름]", "[%]", "전처리, 제목/포스터 특성 추출"],
            ["[학번/이름]", "[%]", "모델 학습, 평가, 시각화"],
            ["[학번/이름]", "[%]", "보고서 작성 및 발표 자료 정리"],
        ],
    )

    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(REPORT_DOCX)


def fmt_int(value: float | int) -> str:
    return f"{int(round(float(value))):,}"


def fmt_float(value: float, digits: int = 3) -> str:
    if pd.isna(value):
        return "-"
    return f"{float(value):.{digits}f}"


def fmt_pct(value: float, digits: int = 1) -> str:
    if pd.isna(value):
        return "-"
    return f"{float(value) * 100:.{digits}f}%"


def format_experiment_model(row: pd.Series) -> str:
    return f"{experiment_label(row['experiment'])} / {model_label(row['model'])}"


def setup_detailed_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    styles["Normal"].font.name = "Apple SD Gothic Neo"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Apple SD Gothic Neo")
    styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Apple SD Gothic Neo")
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
    styles["Normal"].font.size = Pt(10.5)


def add_title_block(doc: Document, stats: dict[str, object]) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(62)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("국내 영화 흥행 여부 예측 및 성공 요인 분석")
    set_run_font(run, size=24, bold=True, color=RGBColor(38, 45, 56))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("영화 제목, 포스터, 메타데이터를 중심으로")
    set_run_font(run, size=14, color=RGBColor(80, 88, 100))

    add_table(
        doc,
        ["항목", "내용"],
        [
            ["교과목명", "데이터마이닝"],
            ["분석 데이터", f"KOBIS/TMDB 국내 영화 {stats['rows']:,}편"],
            ["모델링 기준", "300만 이상 흥행 여부, 관객 수 회귀, 성공도 랭크 분류"],
            ["실행 기준일", "2026년 6월 18일"],
            ["담당 교수명 / 팀명 / 팀원", "[기입]"],
        ],
        widths=[1.7, 4.8],
    )


def add_summary_section(doc: Document, stats: dict[str, object]) -> None:
    best_cls_roc = stats["best_cls_roc"]
    best_cls_f1 = stats["best_cls_f1"]
    best_reg = stats["best_reg_r2"]
    best_rank = stats["best_rank"]
    add_heading(doc, "1. 요약", 1)
    add_paragraph(
        doc,
        f"본 프로젝트는 KOBIS 누적 관객 수와 TMDB 메타데이터를 결합한 국내 영화 {stats['rows']:,}편을 대상으로 "
        "흥행 여부와 관객 규모를 예측한 데이터마이닝 실험이다. 기본 흥행 기준은 누적 관객 수 300만 명 이상이며, "
        "제목 특성과 포스터 이미지 특성이 메타데이터 기반 예측력을 얼마나 보완하는지 ablation study로 비교하였다.",
    )
    add_table(
        doc,
        ["핵심 항목", "실행 결과"],
        [
            ["데이터 범위", f"{stats['year_min']}~{stats['year_max']}년 개봉 영화, 총 {stats['rows']:,}편"],
            ["클래스 분포", f"300만 이상 {stats['success_3m']:,}편, 미만 {stats['rows'] - stats['success_3m']:,}편"],
            ["분류 ROC-AUC 최고", f"{format_experiment_model(best_cls_roc)} / ROC-AUC {best_cls_roc['roc_auc']:.3f}"],
            ["분류 F1 최고", f"{format_experiment_model(best_cls_f1)} / F1 {best_cls_f1['f1']:.3f}"],
            ["회귀 R2 최고", f"{format_experiment_model(best_reg)} / R2 {best_reg['r2']:.3f}, RMSE {best_reg['rmse']:,.0f}명"],
            ["성공도 랭크 최고", f"{format_experiment_model(best_rank)} / macro F1 {best_rank['f1_macro']:.3f}"],
        ],
        widths=[1.7, 4.8],
    )
    add_bullets(
        doc,
        [
            "Random Forest는 300만 기준 분류에서 가장 안정적인 기준 모델이었다.",
            "제목 특성은 ROC-AUC를 일부 보완했지만, 포스터 저수준 특성은 단독 성능 개선 폭이 제한적이었다.",
            "관객 수 회귀의 설명력은 낮아, 실제 관객 수를 직접 맞히기보다 흥행 여부 또는 성공도 등급 예측이 더 실용적이다.",
            "배우·제작비 노트북 분석은 흥행 요인을 해석하는 보조 분석으로 유용하지만, 인과 효과가 아니라 표본 내 연관 신호로 해석해야 한다.",
        ],
    )


def add_data_section(doc: Document, charts: dict[str, Path], stats: dict[str, object]) -> None:
    add_heading(doc, "2. 데이터 구성과 문제 정의", 1)
    add_paragraph(
        doc,
        "원천 데이터는 KOBIS 연도별 박스오피스에서 얻은 국내 영화 관객 수와 TMDB API 기반 영화 상세 정보를 제목 기준으로 매핑한 결과이다. "
        "분석 대상에는 장르, 개봉일, 런타임, 제작비, 감독, 출연진, 제작사, 포스터 URL 및 포스터 이미지 특성이 포함된다.",
    )
    add_table(
        doc,
        ["항목", "값"],
        [
            ["전체 영화 수", f"{stats['rows']:,}편"],
            ["개봉 연도 범위", f"{stats['year_min']}~{stats['year_max']}년"],
            ["포스터 특성 추출 완료", f"{stats['poster_available']:,}편"],
            ["100만 이상", f"{stats['success_1m']:,}편"],
            ["300만 이상", f"{stats['success_3m']:,}편"],
            ["500만 이상", f"{stats['success_5m']:,}편"],
            ["관객 수 평균 / 중앙값", f"{stats['audience_mean']:,}명 / {stats['audience_median']:,}명"],
            ["최대 관객 수", f"{stats['audience_max']:,}명"],
        ],
        widths=[2.0, 4.5],
    )
    add_picture(doc, charts["poster_collage"], width=6.0, caption="그림 1. 데이터에 포함된 주요 흥행작 포스터 예시")
    add_picture(doc, charts["audience_hist"], width=6.3, caption="그림 2. 누적 관객 수 분포")
    add_picture(doc, charts["success_ratio"], width=6.3, caption="그림 3. 흥행 기준별 성공 영화 비율")
    if "rank_distribution" in charts:
        add_picture(doc, charts["rank_distribution"], width=5.5, caption="그림 4. 성공도 랭크 분포")


def add_pipeline_section(doc: Document, charts: dict[str, Path]) -> None:
    add_heading(doc, "3. 전처리와 특성 설계", 1)
    add_paragraph(
        doc,
        "전처리는 관객 수 타깃 생성, 날짜 파생 변수, 장르 원-핫 인코딩, 인물·제작사 빈도 요약, 제목 텍스트 규칙 특성, 포스터 이미지 특성 병합 순서로 수행하였다. "
        "회귀 타깃은 관객 수의 긴 꼬리 분포를 완화하기 위해 log1p 변환을 적용하고, 예측 후 expm1으로 원 단위 관객 수를 복원해 RMSE와 MAE를 계산하였다.",
    )
    add_table(
        doc,
        ["특성 그룹", "주요 변수", "처리 방식"],
        [
            ["메타데이터", "런타임, 제작비, 개봉 연/월/분기, 성수기 여부", "수치 변환, 결측치 중앙값 대체, 제작비 로그 변환"],
            ["장르", "액션, 드라마, 코미디 등 다중 장르", "다중 라벨을 원-핫 변수로 확장"],
            ["감독/배우/제작사", "감독, 출연진, 제작사 목록", "등장 빈도 평균/최댓값, 상위 인물 더미 변수"],
            ["제목", "길이, 단어 수, 숫자/영어/특수문자, 시리즈 키워드", "규칙 기반 텍스트 특성"],
            ["포스터", "밝기, 채도, RGB 평균, 대비, 에지 밀도", "Pillow 기반 저수준 이미지 통계"],
        ],
        widths=[1.15, 2.35, 3.0],
    )
    add_picture(doc, charts["workflow"], width=6.4, caption="그림 5. 전체 분석 파이프라인")


def add_model_design_section(doc: Document) -> None:
    add_heading(doc, "4. 모델 설계", 1)
    add_paragraph(
        doc,
        "실험은 입력 특성 조합을 네 가지로 나누어 진행하였다. metadata는 구조화된 영화 정보만 사용하고, metadata_title은 제목 특성, metadata_poster는 포스터 특성, all은 모든 특성을 포함한다. "
        "이 설계는 제목과 포스터가 기존 메타데이터 대비 실제로 추가 설명력을 갖는지 확인하기 위한 제거 실험이다.",
    )
    add_table(
        doc,
        ["실험명", "입력 특성"],
        [
            ["metadata", "장르, 런타임, 제작비, 개봉 시기, 감독/출연진/제작사 빈도"],
            ["metadata_title", "metadata + 제목 길이, 단어 수, 숫자/영어/특수문자, 키워드"],
            ["metadata_poster", "metadata + 포스터 밝기, 채도, RGB 평균, 대비, 에지 밀도"],
            ["all", "metadata + 제목 특성 + 포스터 특성"],
        ],
        widths=[1.4, 5.1],
    )
    add_table(
        doc,
        ["문제", "모델", "평가지표"],
        [
            ["300만 이상 이진 분류", "Logistic Regression, Decision Tree, Random Forest, DNN MLP", "Accuracy, Precision, Recall, F1, ROC-AUC, Confusion Matrix"],
            ["성공도 랭크 분류", "동일한 분류 모델을 D/C/B/A/S 다중 클래스에 적용", "Accuracy, macro Precision/Recall/F1"],
            ["관객 수 회귀", "Linear, Ridge, Lasso, Random Forest Regressor", "MAE, RMSE, R2, RMSLE"],
        ],
        widths=[1.7, 2.7, 2.1],
    )


def add_classification_section(doc: Document, charts: dict[str, Path], stats: dict[str, object]) -> None:
    cls = stats["classification"]
    top_rows = cls.sort_values("roc_auc", ascending=False).head(8)
    best_roc = stats["best_cls_roc"]
    best_f1 = stats["best_cls_f1"]
    add_heading(doc, "5. 300만 기준 분류 결과", 1)
    add_paragraph(
        doc,
        f"ROC-AUC 기준 최고 모델은 {format_experiment_model(best_roc)}이며 ROC-AUC {best_roc['roc_auc']:.3f}, "
        f"Accuracy {best_roc['accuracy']:.3f}, F1 {best_roc['f1']:.3f}을 기록했다. "
        f"F1 기준 최고 모델은 {format_experiment_model(best_f1)}이며, 흥행작 Recall은 {best_f1['recall']:.3f}이다.",
    )
    add_table(
        doc,
        ["실험", "모델", "Acc", "Prec", "Recall", "F1", "ROC-AUC"],
        classification_rows(top_rows),
        widths=[1.1, 1.25, 0.72, 0.72, 0.78, 0.68, 0.85],
    )
    add_picture(doc, charts["classification"], width=6.2, caption="그림 6. Random Forest 입력 특성 조합별 ROC-AUC")
    add_picture(doc, CONFUSION_MATRIX, width=4.8, caption="그림 7. metadata Random Forest 혼동행렬")
    add_paragraph(
        doc,
        "혼동행렬을 보면 비흥행 영화는 비교적 잘 구분되지만 300만 이상 흥행작 일부가 비흥행으로 예측된다. "
        "따라서 흥행작 발굴이 목적이면 기본 0.5 임계값을 그대로 쓰기보다 Recall을 높이는 threshold tuning 또는 class weight 조정이 필요하다.",
    )


def add_rank_section(doc: Document, charts: dict[str, Path], stats: dict[str, object]) -> None:
    rank = stats["rank"]
    best_rank = stats["best_rank"]
    rank_counts = stats["rank_counts"]
    rows = []
    for _, row in rank.sort_values(["f1_macro", "accuracy"], ascending=False).head(6).iterrows():
        rows.append(
            [
                experiment_label(row["experiment"]),
                model_label(row["model"]),
                fmt_float(row["accuracy"]),
                fmt_float(row["precision_macro"]),
                fmt_float(row["recall_macro"]),
                fmt_float(row["f1_macro"]),
            ]
        )
    add_heading(doc, "6. 성공도 랭크 분류 결과", 1)
    add_paragraph(
        doc,
        "성공도 랭크는 D(100만 미만), C(100만 이상), B(300만 이상), A(500만 이상), S(1000만 이상)로 나눈 다중 클래스 문제이다. "
        f"랭크 분포는 D {rank_counts['D']}편, C {rank_counts['C']}편, B {rank_counts['B']}편, A {rank_counts['A']}편, S {rank_counts['S']}편으로 하위 랭크에 치우쳐 있다.",
    )
    add_table(
        doc,
        ["실험", "모델", "Acc", "Macro P", "Macro R", "Macro F1"],
        rows,
        widths=[1.15, 1.35, 0.8, 0.95, 0.95, 1.0],
    )
    add_paragraph(
        doc,
        f"가장 높은 macro F1은 {format_experiment_model(best_rank)}의 {best_rank['f1_macro']:.3f}이다. "
        "이 값은 이진 분류보다 낮으며, 이는 D/C 중간 구간에 표본이 몰리고 S/A 상위 랭크 표본이 작기 때문이다.",
    )
    for key, caption in [
        ("movie_feature_pca_map", "그림 8. 전체 특성 기반 PCA 2D 분포맵"),
        ("actual_audience_vs_predicted_rank", "그림 9. 실제 관객 수와 예측 성공도 랭크"),
        ("rank_confusion_heatmap", "그림 10. 성공도 랭크 오분류 구조"),
        ("misclassified_movies_lollipop", "그림 11. 관객 수가 큰 오분류 영화"),
    ]:
        if key in charts:
            add_picture(doc, charts[key], width=6.2, caption=caption)


def add_regression_section(doc: Document, charts: dict[str, Path], stats: dict[str, object]) -> None:
    reg = stats["regression"]
    best_r2 = stats["best_reg_r2"]
    rows = []
    for _, row in reg.sort_values("r2", ascending=False).head(6).iterrows():
        rows.append(
            [
                experiment_label(row["experiment"]),
                model_label(row["model"]),
                fmt_int(row["mae"]),
                fmt_int(row["rmse"]),
                fmt_float(row["r2"]),
                fmt_float(row["rmsle"]),
            ]
        )
    add_heading(doc, "7. 관객 수 회귀 결과", 1)
    add_paragraph(
        doc,
        f"관객 수 회귀에서 R2 기준 최고 모델은 {format_experiment_model(best_r2)}이며 R2 {best_r2['r2']:.3f}, "
        f"RMSE {best_r2['rmse']:,.0f}명이다. 절대 관객 수 예측은 블록버스터의 긴 꼬리와 외부 변수 누락 때문에 설명력이 제한적이었다.",
    )
    add_table(
        doc,
        ["실험", "모델", "MAE", "RMSE", "R2", "RMSLE"],
        rows,
        widths=[1.15, 1.45, 1.1, 1.15, 0.75, 0.8],
    )
    add_picture(doc, charts["regression"], width=6.3, caption="그림 12. 관객 수 회귀 모델 RMSE 비교")


def add_factor_section(doc: Document, charts: dict[str, Path]) -> None:
    add_heading(doc, "8. 성공 요인 해석", 1)
    add_paragraph(
        doc,
        "Random Forest 변수 중요도는 런타임, 제작비, 출연진·감독 빈도, 개봉 연도, 포스터 색상 특성 등이 함께 작용함을 보여준다. "
        "단일 변수가 흥행을 결정하기보다 작품 규모, 인물 네트워크, 장르, 개봉 시기, 시각적 인상이 결합되는 구조로 해석하는 것이 적절하다.",
    )
    add_picture(doc, charts["importance"], width=6.3, caption="그림 13. 전체 특성 Random Forest 변수 중요도 Top 12")
    add_picture(doc, charts["genre_avg"], width=6.3, caption="그림 14. 장르별 평균 관객 수 Top 10")
    for key, caption in [
        ("genre_rank_heatmap", "그림 15. 장르별 성공도 랭크 구성비"),
        ("month_rank_heatmap", "그림 16. 개봉 월별 성공도 랭크 구성비"),
        ("poster_success", "그림 17. 흥행 여부별 포스터 색상 특성 평균"),
        ("poster_brightness_saturation_scatter", "그림 18. 포스터 밝기/채도와 성공도 랭크"),
    ]:
        if key in charts:
            add_picture(doc, charts[key], width=6.2, caption=caption)


def add_notebook_analysis_section(doc: Document, charts: dict[str, Path], stats: dict[str, object]) -> None:
    actor_lasso = stats["actor_lasso"]
    actor_hit = stats["actor_hit"]
    budget = stats["budget_cluster"]
    add_heading(doc, "9. 추가 노트북 분석", 1)
    add_paragraph(
        doc,
        "별도 노트북에서 수행한 배우 연관 신호와 제작비-관객 수 군집화도 보고서에 포함하였다. "
        "이 분석은 모델 성능 향상보다는 결과 해석과 후속 분석 방향을 잡기 위한 탐색적 분석이다.",
    )

    add_heading(doc, "9.1 배우별 관객 수 연관 신호", 2)
    lasso_metrics = actor_lasso["metrics"]
    add_table(
        doc,
        ["모델", "MAE", "RMSE", "R2"],
        [[row["model"], fmt_int(row["mae"]), fmt_int(row["rmse"]), fmt_float(row["r2"])] for _, row in lasso_metrics.iterrows()],
        widths=[2.0, 1.45, 1.45, 0.8],
    )
    lasso_rows = []
    for _, row in actor_lasso["top_positive"].iterrows():
        lasso_rows.append(
            [
                row["actor"],
                fmt_int(row["n_movies"]),
                f"{row['estimated_delta_million']:.2f}",
                f"{row['mean_audience_million']:.2f}",
                fmt_pct(row["hit_rate_5m"]),
            ]
        )
    add_table(
        doc,
        ["배우", "출연편수", "추정 차이(백만)", "평균 관객(백만)", "500만 비율"],
        lasso_rows,
        widths=[1.55, 0.8, 1.25, 1.25, 1.0],
    )
    add_picture(doc, charts["actor_lasso_signal"], width=6.2, caption="그림 19. Lasso 배우별 관객 수 연관 신호")
    add_paragraph(
        doc,
        "Lasso 계수는 다른 배우 신호가 동시에 들어간 상태의 연관 신호다. 출연 편수가 적은 배우는 우연성이 크므로 평균 관객 수와 출연편수를 함께 확인해야 한다.",
    )

    add_heading(doc, "9.2 배우별 500만 돌파 확률", 2)
    hit_metrics = actor_hit["metrics"]
    add_table(
        doc,
        ["모델", "Acc", "Prec", "Recall", "F1", "ROC-AUC"],
        [
            [
                row["model"],
                fmt_float(row["accuracy"]),
                fmt_float(row["precision"]),
                fmt_float(row["recall"]),
                fmt_float(row["f1"]),
                fmt_float(row["roc_auc"]),
            ]
            for _, row in hit_metrics.iterrows()
        ],
        widths=[1.8, 0.75, 0.8, 0.85, 0.75, 0.9],
    )
    add_paragraph(
        doc,
        "아래 표는 확률 상위 배우와 함께, 데이터 내 출연편수와 평균 관객 수를 기준으로 상위 인지도군보다 한 단계 낮게 잡힌 중저인지도 후보를 별도로 보여준다. "
        "인지도는 외부 검색량이나 설문 지표가 아니라 본 데이터 내부의 노출도 proxy이므로 보조 해석으로만 사용한다.",
    )
    prob_rows = []
    for _, row in actor_hit["top_probability"].iterrows():
        prob_rows.append(
            [
                row["actor"],
                fmt_int(row["n_movies"]),
                f"{row['single_actor_hit_prob_pct']:.1f}%",
                fmt_pct(row["hit_rate_5m"]),
                f"{row['mean_audience_million']:.2f}",
            ]
        )
    add_table(
        doc,
        ["배우", "출연편수", "단독 확률", "실제 500만 비율", "평균 관객(백만)"],
        prob_rows,
        widths=[1.55, 0.8, 1.0, 1.2, 1.25],
    )
    lower_profile_rows = []
    for _, row in actor_hit["lower_profile_probability"].iterrows():
        lower_profile_rows.append(
            [
                row["actor"],
                fmt_int(row["n_movies"]),
                f"{row['single_actor_hit_prob_pct']:.1f}%",
                fmt_pct(row["hit_rate_5m"]),
                f"{row['mean_audience_million']:.2f}",
            ]
        )
    add_table(
        doc,
        ["중저인지도 후보", "출연편수", "단독 확률", "실제 500만 비율", "평균 관객(백만)"],
        lower_profile_rows,
        widths=[1.55, 0.8, 1.0, 1.2, 1.25],
    )
    add_picture(doc, charts["actor_hit_probability"], width=6.2, caption="그림 20. 배우 신호 기반 500만 돌파 확률 후보")
    add_paragraph(
        doc,
        "단독 확률은 특정 배우 더미만 켠 가상 입력에 대한 모델 출력이다. 실제 영화의 장르, 제작비, 감독, 개봉 시기를 반영한 확률이 아니므로 캐스팅 효과를 단정하면 안 된다.",
    )

    add_heading(doc, "9.3 제작비-관객 수 군집", 2)
    budget_rows = []
    for _, row in budget["summary"].iterrows():
        budget_rows.append(
            [
                row["cluster_label"],
                fmt_int(row["n_movies"]),
                f"{row['median_budget_million']:.2f}",
                f"{row['median_audience_million']:.2f}",
                fmt_pct(row["hit_rate_3m"]),
                f"{row['median_audience_per_budget_million']:.2f}",
            ]
        )
    add_table(
        doc,
        ["군집", "영화 수", "중앙 제작비", "중앙 관객", "300만 비율", "관객/제작비"],
        budget_rows,
        widths=[1.35, 0.75, 1.0, 1.0, 0.95, 1.0],
    )
    add_picture(doc, charts["budget_audience_cluster"], width=6.3, caption="그림 21. 제작비-관객 수 기반 영화 군집")
    add_paragraph(
        doc,
        f"제작비가 0이 아니고 관객 수가 있는 {budget['n_movies']:,}편을 대상으로 KMeans {budget['selected_k']}개 군집을 구성했다. "
        "TMDB 제작비는 결측과 통화 기준 문제가 있을 수 있으므로 절대 손익분기점이 아니라 상대적 포지셔닝으로 해석해야 한다.",
    )


def add_limitations_section(doc: Document) -> None:
    add_heading(doc, "10. 한계점과 개선 방향", 1)
    add_table(
        doc,
        ["한계점", "영향", "개선 방향"],
        [
            ["표본 수 654편", "상위 흥행작과 최신 영화 표본이 작음", "연도별 데이터 확장, 해외/OTT 성과 추가"],
            ["제목 기반 KOBIS/TMDB 매핑", "동명 영화 또는 개봉연도 차이로 매칭 오류 가능", "영화 코드, 감독, 개봉일을 함께 쓰는 매칭 검증"],
            ["마케팅/배급 변수 부재", "관객 수 회귀 설명력 제한", "스크린 수, 배급사, 광고비, 예매량, SNS 지표 수집"],
            ["포스터 저수준 특성 중심", "이미지 의미 정보 반영 부족", "CLIP/CNN 임베딩과 포스터 텍스트 OCR 추가"],
            ["클래스 불균형", "흥행작 Recall 저하", "임계값 조정, class weight, 재표본화, 교차검증 적용"],
        ],
        widths=[1.45, 2.0, 3.05],
    )
    add_paragraph(
        doc,
        "결론적으로 본 프로젝트는 데이터 수집부터 전처리, 특성 설계, 모델 학습, 평가, 해석까지 일관된 파이프라인을 완성했다. "
        "현재 단계에서는 흥행 여부 분류가 관객 수 회귀보다 실용적이며, 실제 적용을 위해서는 외부 변수 추가와 검증 체계 강화가 가장 중요하다.",
    )


def add_references_section(doc: Document) -> None:
    add_heading(doc, "11. 참고문헌 및 기여도", 1)
    add_table(
        doc,
        ["구분", "자료"],
        [
            ["KOBIS", "영화관입장권통합전산망 박스오피스 및 연도별 통계"],
            ["TMDB", "The Movie Database API: 영화 상세 정보, 장르, 포스터, 출연진 데이터"],
            ["scikit-learn", "Logistic Regression, Decision Tree, Random Forest, MLP, KMeans, PCA"],
            ["Pillow / matplotlib", "포스터 이미지 처리와 보고서용 시각화 생성"],
            ["프로젝트 코드", "src/preprocessing.py, src/train_classification.py, src/train_regression.py, notebooks/*.ipynb"],
        ],
        widths=[1.4, 5.1],
    )
    add_paragraph(doc, "팀원별 기여도는 실제 팀 구성에 맞춰 아래 표를 수정하면 된다.")
    add_table(
        doc,
        ["팀원", "기여도", "역할"],
        [
            ["[학번/이름]", "[%]", "데이터 수집 및 KOBIS/TMDB 매핑"],
            ["[학번/이름]", "[%]", "전처리, 제목/포스터 특성 추출"],
            ["[학번/이름]", "[%]", "모델 학습, 평가, 시각화"],
            ["[학번/이름]", "[%]", "보고서 작성 및 발표 자료 정리"],
        ],
        widths=[1.5, 1.0, 4.0],
    )


def build_detailed_docx(frame: pd.DataFrame, charts: dict[str, Path], stats: dict[str, object]) -> None:
    doc = Document()
    setup_detailed_styles(doc)
    add_title_block(doc, stats)
    add_picture(doc, charts["poster_collage"], width=6.0)
    doc.add_page_break()

    add_summary_section(doc, stats)
    doc.add_page_break()
    add_data_section(doc, charts, stats)
    doc.add_page_break()
    add_pipeline_section(doc, charts)
    add_model_design_section(doc)
    doc.add_page_break()
    add_classification_section(doc, charts, stats)
    doc.add_page_break()
    add_rank_section(doc, charts, stats)
    doc.add_page_break()
    add_regression_section(doc, charts, stats)
    doc.add_page_break()
    add_factor_section(doc, charts)
    doc.add_page_break()
    add_notebook_analysis_section(doc, charts, stats)
    doc.add_page_break()
    add_limitations_section(doc)
    doc.add_page_break()
    add_references_section(doc)

    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(REPORT_DOCX)


def markdown_report(stats: dict[str, object]) -> str:
    best_cls = stats["best_cls_roc"]
    best_cls_f1 = stats["best_cls_f1"]
    best_reg = stats["best_reg_r2"]
    best_rank = stats["best_rank"]
    return f"""# 국내 영화 흥행 여부 예측 및 성공 요인 분석

영화 제목, 포스터, 메타데이터를 중심으로

## 표지 정보

| 항목 | 내용 |
| --- | --- |
| 교과목명 | 데이터마이닝 |
| 담당 교수명 | [기입] |
| 팀명 | [기입] |
| 팀원 학번 / 이름 | [기입] |
| 제출일 | 2026년 6월 |

## 1. 프로젝트 개요

본 프로젝트는 KOBIS 박스오피스 관객 수와 TMDB 영화 메타데이터를 결합하여 국내 영화의 흥행 여부를 예측하고, 흥행에 영향을 주는 요인을 분석하는 것을 목표로 한다. 분석 대상은 총 {stats['rows']}편이며, 기본 흥행 기준은 누적 관객 수 300만 명 이상으로 설정하였다.

## 2. 데이터 설명

| 항목 | 값 |
| --- | --- |
| 전체 영화 수 | {stats['rows']:,}편 |
| 개봉 연도 범위 | {stats['year_min']}~{stats['year_max']}년 |
| 포스터 특성 추출 완료 | {stats['poster_available']:,}편 |
| 100만 이상 영화 | {stats['success_1m']:,}편 |
| 300만 이상 영화 | {stats['success_3m']:,}편 |
| 500만 이상 영화 | {stats['success_5m']:,}편 |
| 관객 수 평균 | {stats['audience_mean']:,}명 |
| 관객 수 중앙값 | {stats['audience_median']:,}명 |

## 3. 전처리 및 특성 설계

메타데이터는 런타임, 제작비, 개봉 연도와 월, 계절 변수, 장르 원-핫 인코딩, 감독·출연진·제작사 빈도 변수로 구성하였다. 제목 특성은 길이, 단어 수, 숫자/영어/특수문자 포함 여부, 장르성 키워드 여부로 설계하였다. 포스터 특성은 평균 밝기, 평균 채도, RGB 평균, 대비, 에지 밀도 등 이미지의 저수준 시각 특성으로 추출하였다.

## 4. 모델 설계

분류 모델은 Logistic Regression, Decision Tree, Random Forest, DNN MLP를 비교하였다. 회귀 모델은 Linear Regression, Ridge, Lasso, Random Forest Regressor를 비교하였다. 입력 특성 조합은 metadata, metadata_title, metadata_poster, all의 네 가지 실험으로 나누어 ablation study를 수행하였다.

## 5. 분류 결과

ROC-AUC 기준 최상위 모델은 {experiment_label(best_cls['experiment'])} / {model_label(best_cls['model'])}이며 ROC-AUC는 {best_cls['roc_auc']:.3f}, Accuracy는 {best_cls['accuracy']:.3f}, F1은 {best_cls['f1']:.3f}이다. F1 기준 최상위 모델은 {experiment_label(best_cls_f1['experiment'])} / {model_label(best_cls_f1['model'])}이며 F1은 {best_cls_f1['f1']:.3f}이다.

## 6. 회귀 결과

R2 기준 최상위 회귀 모델은 {experiment_label(best_reg['experiment'])} / {model_label(best_reg['model'])}이며 R2는 {best_reg['r2']:.3f}, RMSE는 {best_reg['rmse']:,.0f}명이다. 관객 수 분포가 블록버스터 영화에 크게 치우쳐 있어 회귀 예측의 설명력은 제한적이었다.

## 7. 성공도 랭크 및 추가 노트북 분석

성공도 랭크 분류의 최상위 모델은 {experiment_label(best_rank['experiment'])} / {model_label(best_rank['model'])}이며 macro F1은 {best_rank['f1_macro']:.3f}이다. 추가 노트북 분석으로 배우별 관객 수 연관 신호, 배우별 500만 돌파 확률, 제작비-관객 수 군집화를 실행해 상세 DOCX에 그림과 표로 반영하였다.

## 8. 주요 분석 결과

- Random Forest는 분류 문제에서 가장 안정적인 기준 모델로 확인되었다.
- 제목 특성은 ROC-AUC를 소폭 높였으나 전체 성능 개선 폭은 크지 않았다.
- 포스터 특성은 변수 중요도 상위권 일부에 포함되지만 단독으로 큰 성능 향상을 만들지는 못했다.
- 흥행작 클래스가 적어 Accuracy만 보면 성능이 과대평가될 수 있으므로 F1, Recall, ROC-AUC를 함께 해석해야 한다.

## 9. 한계점 및 향후 개선 방향

데이터 수가 654편으로 제한적이고 일부 제목 매칭 오류 가능성이 있다. 또한 제작비 결측, 마케팅비 부재, 개봉 직전 예매량과 SNS 반응 같은 외부 변수가 포함되지 않았다. 향후에는 영화 코드 기반 매칭 고도화, CLIP/CNN 기반 포스터 임베딩, 교차검증, 클래스 불균형 보정, 임계값 조정 등을 적용할 수 있다.

## 10. 참고문헌

- KOBIS 영화관입장권통합전산망 박스오피스 및 연도별 통계
- TMDB API 영화 상세 정보, 장르, 포스터, 출연진 데이터
- scikit-learn 공식 문서
- Pillow 이미지 처리 라이브러리

## 11. 팀원별 기여도

| 팀원 | 기여도 | 역할 |
| --- | --- | --- |
| [학번/이름] | [%] | 데이터 수집 및 KOBIS/TMDB 매핑 |
| [학번/이름] | [%] | 전처리, 제목/포스터 특성 추출 |
| [학번/이름] | [%] | 모델 학습, 평가, 시각화 |
| [학번/이름] | [%] | 보고서 작성 및 발표 자료 정리 |
"""


def main() -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    frame = pd.read_csv(DATA_FILE, encoding="utf-8-sig")
    charts = create_charts(frame)
    stats = table_data(frame)
    build_detailed_docx(frame, charts, stats)
    REPORT_MD.write_text(markdown_report(stats), encoding="utf-8")
    print(f"docx: {REPORT_DOCX}")
    print(f"markdown: {REPORT_MD}")
    print(f"figures: {FIGURE_DIR}")


if __name__ == "__main__":
    main()
